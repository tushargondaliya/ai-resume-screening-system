"""
Resume Parser - Extract text from PDF and DOCX resume files.
Uses PyPDF2 for PDFs and python-docx for DOCX files.
"""
import os

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None


def parse_resume(file_path):
    """Parse a resume file and extract text."""
    if not os.path.exists(file_path):
        return ''
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    else:
        return ''


def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    if PdfReader is None:
        return '[PDF parsing unavailable - install PyPDF2]'
    
    try:
        reader = PdfReader(file_path)
        text = ''
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
        return text.strip()
    except Exception as e:
        return f'[Error reading PDF: {str(e)}]'


def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    if docx is None:
        return '[DOCX parsing unavailable - install python-docx]'
    
    try:
        doc = docx.Document(file_path)
        text = ''
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + '\n'
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += ' | '.join(row_text) + '\n'
        
        return text.strip()
    except Exception as e:
        return f'[Error reading DOCX: {str(e)}]'


def extract_text_from_txt(file_path):
    """Extract text from a plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read().strip()
    except Exception as e:
        return f'[Error reading TXT: {str(e)}]'
